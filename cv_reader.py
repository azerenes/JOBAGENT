import re
import os

STOPWORDS = set("""a acaba adam ayni ayrca akar alti altin ama ancak arda ardindan arasinda ari artik
asla at bazi belki ben beni benim beri bey bina binler biz bizim bosa boyle bu buda buyuk buna bunca
bunlar buralari cay cok cunku cunku daha dahi da de dek demek de mi daha diye diger diye dun eger eksi
en er evet evet eyle ey ne fazla fakat gibi goz gore goruldugu gore gorev guzel hakkinda hala hangi
hani hatta hemen hemhenuz her herkes hic icin icin ilave ileride ile ilgili ilk imkanimi iste iste
kadar karsilik karsi kendisi kim kimse lazim lakin macar madem magar maalesef mali meger mesele mecburen
mi mu mukemmel muste neden nerede neyse nice nokta o olan olsun on ondan onlar onu oysa oysaki ozellikle
ozur pek peki pekala raga razi sag su suan sure sen seni sira sira son sonra sonraki soz su sizin taa
tamam tamamen tarih tek tesekkur teyze ugrasam uslu uzun vardir varken varveyahut veya yani yapiskan
yerine yeter yine yok zaten zaten""".split())


def normalize_turkish(s):
    tr = str.maketrans('İışğüöçİIŞĞÜÖÇ', 'iisguociisguoc')
    return s.translate(tr).lower()


def extract_text(path):
    """PDF / DOCX / TXT / DOC dosyasindan duz metin cikarir."""
    ext = (path or '').lower().rsplit('.', 1)[-1]
    if ext == 'pdf':
        from pypdf import PdfReader
        r = PdfReader(path)
        return '\n'.join((p.extract_text() or '') for p in r.pages)
    if ext == 'docx':
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for tb in d.tables:
            for row in tb.rows:
                parts += [c.text for c in row.cells]
        return '\n'.join(parts)
    if ext in ('txt', 'md', 'rtf'):
        with open(path, encoding='utf-8', errors='replace') as f:
            return f.read()
    # bilinmeyen: metin olarak dene
    try:
        with open(path, encoding='utf-8', errors='replace') as f:
            return f.read()
    except Exception:
        return ''


TECH_PREFIXES = ('zayif', 'enerji', 'pano', 'kumanda', 'kablo', 'elektrik', 'elektron', 'otomas', 'montaj',
                 'bakim', 'saha', 'uretim', 'kalite', 'ges', 'guc', 'isik', 'aydin', 'trafo', 'motor', 'sarj',
                 'scada', 'devre', 'sistem', 'teknis', 'kontrol', 'endustri', 'test', 'cihaz', 'makin', 'robot')


def extract_keywords(text):
    """Metinden anlamli anahtar kelimeleri cikarir (tekil tokenlar + kisa teknik ifadeler)."""
    tl = normalize_turkish(text)
    tl = re.sub(r'[^a-z0-9\s/+-]', ' ', tl)
    tokens = [w for w in re.split(r'\s+', tl) if w.strip() and w not in STOPWORDS and len(w) >= 3]
    out = set()
    for w in tokens:
        if any(k in w for k in TECH_PREFIXES):
            out.add(w)
    for i in range(len(tokens) - 1):
        w1, w2 = tokens[i], tokens[i + 1]
        if w1 in TECH_PREFIXES or any(w1.startswith(k) for k in TECH_PREFIXES):
            out.add(f'{w1} {w2}')
    return sorted(out)
